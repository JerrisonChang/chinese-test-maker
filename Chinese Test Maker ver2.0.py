#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Mar  5 10:46:14 2019

@author: jerrisonchang
"""
import pandas as pd
import random
import docx
from docx.shared import Pt
from docx.shared import Inches


def write_new_paragraph(document, string, font_size = None, font_name = None, bold = None, underline = None):
    '''
    
    '''
    new_paragraph = document.add_paragraph(string)
    new_paragraph.runs[0].font.size = Pt(font_size)
    new_paragraph.runs[0].font.name = font_name
    new_paragraph.runs[0].bold = bold
    new_paragraph.runs[0].underline = underline
    
    if bold or underline == True:
        run = new_paragraph.add_run('')
        run.bold = False
        run.underline = False
        
    return new_paragraph

def thorough_shuffle(list_old):
    '''
    take in a list and return shuffled sentence, joining with ' / '
    It also checks if the new list is too similar to the original sentence.
    '''
    n = len(list_old)
    list_new = list_old.copy()
    
    criteria = 3 #number of the words order intact
    while n >= criteria:
        n = len(list_old)
        random.shuffle(list_new)
        for i in range(len(list_old)):
            if list_new[i] == list_old[i]:
                n =- 1
    final_sentence_ch = " / ".join(list_new)
    
    return final_sentence_ch


def rearrange_char(data, n):
    '''
    read row n from excel and shuffle the chinese part,
    then add English explanation at the bottom
    '''
    ##Step 1: Create a list from the sentence and shuffle it
    sentence_ch = data.loc[data.index[n],'Chinese Question'].strip()
    sentence_eng = data.loc[data.index[n],'English Translation/ Target word'].strip()
    sentence_ch_list = sentence_ch.split(' ')
    final_sentence_ch = thorough_shuffle(sentence_ch_list)

    ##Step 2: Write down Chinese sentence
    new_paragraph = write_new_paragraph(document, final_sentence_ch, 14, 'KaiTi SC')
    new_paragraph.paragraph_format.space_after = Pt(30)
    
    ##Step 3: Write down English traslation    
    run = new_paragraph.add_run('\n({})'.format(sentence_eng))
    run.font.size = Pt(11)
    
    return None

def rearrange_char_set(qbank, document):
    '''
    generate the problem set from sheet_number
    (sheet_number starts from 0)
    and write into document
    '''
    data = qbank

    ##Step 1
    stem = 'Part I: Rearrange the characters'
    write_new_paragraph(document, stem, 14, bold = True, underline = True)

    ##Step 2: Print out the questions
    for i in range(len(data.index)):
        rearrange_char(data,i)
    
    return None

def dots_in(string):
    '''
    check is there dots in the string and return a tuple
    '''
    if '...' in string:
        result = (True, '...')
    elif '…' in string:
        result = (True, '…')
    else:
        result = (False, None)
    return result

def create_word_pool(data):
    '''
    create word pool to choose from
    '''
    ##Step 1: Fetch the words
    word_pool = []
    for i in range(len(data.index)):
        value = data.loc[data.index[i],'English Translation/ Target word'].strip()
        if ' ' in value:
            sub_list = value.split(' ')
            word_pool.extend(sub_list)
        elif '/' in value:
            None
        else:
            word_pool.append(value)
    random.shuffle(word_pool)

    ##Step 2: Chunk the list into pieces
    list_len = len(word_pool)
    if list_len >= 10 & list_len%6 != 1:
        sn = 6 #max number of words per line
    else:
        sn = 5
    
    word_pool_sets = []
    for i in range(0,list_len,sn):
        word_pool_sets.append("、".join(word_pool[i:i+sn]))
    final_word_pool = "\n".join(word_pool_sets)

    ##Step 3: Write them down
    new_paragraph = write_new_paragraph(document, final_word_pool, 14, 'KaiTi SC')
    new_paragraph.paragraph_format.left_indent = Inches(0.75)

    return None
    
def create_blank(data,n):
    '''
    create blanded sentence
    n stands each question
    '''
    ##Step 1: Set up
    original = data.loc[data.index[n],'Chinese Question'].strip()
    word_string = data.loc[data.index[n],'English Translation/ Target word'].strip()
    
    new_sentence = original

    blank = '__________'
    sub_list = []

    if '/' in word_string:
        sub_list = word_string.split('/')
        random.shuffle(sub_list)
        new_sentence = new_sentence.replace(word_string[0]+' ',blank), '({})'.format(' / '.join(sub_list))
    else:

    ##Step 2: check multiple words      
        if ' ' in word_string:
            sub_list = word_string.split(' ')
        else:
            sub_list.append(word_string)
        
    ##Step 3: Check grammer
        for i in sub_list:
            if dots_in(i)[0]:
                i = i.strip(dots_in(i)[1])
                ssub_list = i.split(dots_in(i)[1])
                for j in ssub_list:
                    new_sentence = new_sentence.replace(j,blank)
            else:
                new_sentence = new_sentence.replace(i,blank)

    
    ##Step 4: Write down sentences with blank
    write_new_paragraph(document,new_sentence,14,'Kaiti SC')
    
    return None

def fill_in_the_blank_set(qbank , document):
    '''
    generate the problem set from sheet_number
    (sheet_number starts from 0)
    and write into document
    '''
    data = qbank
    
    ##Step 1: Set up
    stem = 'Part II: Fill in the blank'
    write_new_paragraph(document, stem, 14, bold = True, underline = True)

    ##Step 2: Print out word pool
    create_word_pool(data)

    ##Step 3: Print out questions
    for i in range(len(data.index)):
        create_blank(data,i)
        
    return None

if __name__ == '__main__':
    document = docx.Document()

    read_name = 'Chinese question bank.xlsx'
    controller = pd.read_excel(read_name , sheet_name = 'Controller' , header = None, index_col = 0)
    
    #Retrieve the parameters
    class_level = controller.loc['Level',1]
    lesson = controller.loc['Lesson',1]
    purpose = controller.loc['Type',1]
    qyear = controller.loc['Year',1]
    
    document_title = '{0} Lesson {1} {2}'.format(class_level,lesson,purpose)
    document_name = document_title + ' {}.docx'.format(qyear)
    
    qbank = pd.read_excel(read_name , sheet_name = class_level , header = 0 , index_col = None)

    #Filter the data for different Question Type    
    qbank_rearrange = qbank[(qbank['Lesson']== lesson) & (qbank['Type']== purpose) & (qbank['Question Type'].map(lambda x: x.startswith('1')))]
    qbank_fillin = qbank[(qbank['Lesson']== lesson) & (qbank['Type']== purpose) & (qbank['Question Type'].map(lambda x: x.startswith('2')))]
    
    write_new_paragraph(document, document_title, 18, bold = True, underline = True)

    rearrange_char_set(qbank_rearrange, document)
    fill_in_the_blank_set(qbank_fillin,document)    
    
    document.save(document_name)
